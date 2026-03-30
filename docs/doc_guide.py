#!/usr/bin/env python3
"""Generate User Guide DOCX for SynGT C++ application."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(18)
    elif level == 2:
        hs.font.size = Pt(16)
    elif level == 3:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(14)

# ── Title ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Руководство пользователя\nсистемы SynGT (C++)')
run.bold = True
run.font.size = Pt(22)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Краткий гид по приложению')
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

doc.add_paragraph()

# ── TOC ──
doc.add_heading('Содержание', level=1)
toc_items = [
    '1. Общие сведения',
    '2. Запуск и начальный экран',
    '3. Интерфейс графического приложения',
    '  3.1. Панель редактора грамматики',
    '  3.2. Область визуализации диаграммы',
    '  3.3. Список нетерминалов',
    '  3.4. Редактор правила',
    '  3.5. Панель операций',
    '  3.6. Консоль вывода',
    '4. Меню приложения',
    '  4.1. File — работа с файлами',
    '  4.2. Edit — отмена и повтор',
    '  4.3. Grammar — операции с грамматикой',
    '  4.4. Help — справка',
    '5. Синтаксис RBNF-грамматик',
    '  5.1. Элементы грамматики',
    '  5.2. Операторы',
    '  5.3. Примеры грамматик',
    '6. Работа с диаграммой',
    '  6.1. Навигация и масштабирование',
    '  6.2. Выделение объектов',
    '  6.3. Контекстное меню диаграммы',
    '7. Преобразования грамматик',
    '8. Анализ грамматик',
    '9. Вспомогательные функции',
    '10. Горячие клавиши',
    '11. Консольное приложение syngt_cli',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ── 1. Общие сведения ──
doc.add_heading('1. Общие сведения', level=1)
doc.add_paragraph(
    'SynGT (Syntax Graph Tool) — инструментальная система для построения '
    'и эквивалентных преобразований синтаксических диаграмм по RBNF-грамматикам. '
    'Система позволяет визуализировать формальные грамматики в виде диаграмм Вирта '
    '(railroad diagrams), редактировать их и применять различные преобразования.'
)
doc.add_paragraph(
    'Система состоит из трёх компонентов:'
)
items = [
    ('libsyngt', 'библиотека с алгоритмами разбора, визуализации, преобразования и анализа грамматик.'),
    ('syngt_gui', 'графическое приложение с интерфейсом на базе Dear ImGui.'),
    ('syngt_cli', 'консольное приложение для пакетной обработки грамматик из командной строки.'),
]
for name, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    p.add_run(' — ' + desc)

# ── 2. Запуск ──
doc.add_heading('2. Запуск и начальный экран', level=1)
doc.add_paragraph(
    'После запуска syngt_gui открывается главное окно приложения. '
    'По умолчанию создаётся пустая грамматика с начальным нетерминалом S и '
    'пустым (epsilon) правилом. На диаграмме отображается стрелка от начала к концу, '
    'обозначающая пустое правило.'
)
doc.add_paragraph(
    'Для начала работы можно:'
)
bullets = [
    'Ввести грамматику вручную в текстовый редактор (левая панель) и нажать Parse (F5).',
    'Загрузить грамматику из .grm файла через File → Open (Ctrl+O).',
    'Импортировать грамматику из формата GEdit через File → Import from GEdit.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ── 3. Интерфейс ──
doc.add_heading('3. Интерфейс графического приложения', level=1)
doc.add_paragraph(
    'Главное окно разделено на несколько панелей. Ниже описано назначение каждой.'
)

doc.add_heading('3.1. Панель редактора грамматики (Grammar Editor)', level=2)
doc.add_paragraph(
    'Расположена в левой части окна (вкладка «Grammar Editor»). Содержит текстовое поле, '
    'в которое вводится грамматика в формате RBNF. Каждое правило записывается в формате:'
)
p = doc.add_paragraph()
run = p.add_run('    Нетерминал : Правило.')
run.font.name = 'Courier New'
run.font.size = Pt(12)

doc.add_paragraph(
    'Грамматика завершается маркером EOGram! (End Of Grammar). '
    'Под текстовым полем расположены кнопки:'
)
bullets = [
    'Parse (F5) — разобрать грамматику и построить диаграммы для всех нетерминалов.',
    'Clear — очистить текст грамматики и сбросить все данные.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('3.2. Область визуализации диаграммы (Syntax Diagram)', level=2)
doc.add_paragraph(
    'Расположена в центральной части окна (вкладка «Syntax Diagram»). Отображает синтаксическую '
    'диаграмму для выбранного нетерминала. Элементы диаграммы:'
)
table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Элемент', 'Описание']
data = [
    ('Треугольник (▶)', 'Начало и конец диаграммы'),
    ('Овал', 'Терминальный символ (например, \'if\', \'+\')'),
    ('Прямоугольник', 'Нетерминальный символ (ссылка на другое правило)'),
    ('Пунктирный прямоугольник', 'Макроопределение (@macro)'),
    ('Точка разветвления', 'Точка начала/конца альтернативы'),
    ('Стрелка', 'Связь между элементами (направление обработки)'),
]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
for i, (el, desc) in enumerate(data):
    table.rows[i+1].cells[0].text = el
    table.rows[i+1].cells[1].text = desc

doc.add_heading('3.3. Список нетерминалов', level=2)
doc.add_paragraph(
    'Расположен в верхней части правой панели на вкладке «Syntax Diagram». '
    'Отображает все нетерминалы текущей грамматики. Клик по нетерминалу переключает '
    'визуализацию на его диаграмму. Кнопка «Delete NT» удаляет выбранный нетерминал.'
)

doc.add_heading('3.4. Редактор правила', level=2)
doc.add_paragraph(
    'Расположен под списком нетерминалов. Содержит текстовое поле с правилом '
    'выбранного нетерминала. Можно отредактировать правило вручную и нажать «Build (Apply Changes)» '
    'для применения изменений и перестроения диаграммы.'
)

doc.add_heading('3.5. Панель операций (Operations)', level=2)
doc.add_paragraph(
    'Расположена в правой части окна. Содержит кнопки для применения преобразований '
    'и анализа грамматики (подробнее в разделах 7–9).'
)

doc.add_heading('3.6. Консоль вывода (Output)', level=2)
doc.add_paragraph(
    'Расположена в нижней части окна. Отображает сообщения о результатах операций: '
    'ошибки разбора, результаты проверки LL(1), информацию о преобразованиях и т.д.'
)

doc.add_page_break()

# ── 4. Меню ──
doc.add_heading('4. Меню приложения', level=1)

doc.add_heading('4.1. File — работа с файлами', level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['Пункт меню', 'Клавиша', 'Описание']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
data = [
    ('New', 'Ctrl+N', 'Создать новую пустую грамматику (S : eps)'),
    ('Open...', 'Ctrl+O', 'Загрузить грамматику из .grm файла'),
    ('Save', 'Ctrl+S', 'Сохранить текущую грамматику'),
    ('Save As...', '—', 'Сохранить грамматику в новый файл'),
    ('Import from GEdit...', '—', 'Импортировать грамматику из формата GEdit'),
]
for i, (menu, key, desc) in enumerate(data):
    table.rows[i+1].cells[0].text = menu
    table.rows[i+1].cells[1].text = key
    table.rows[i+1].cells[2].text = desc

doc.add_heading('4.2. Edit — отмена и повтор', level=2)
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['Пункт меню', 'Клавиша', 'Описание']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
data = [
    ('Undo', 'Ctrl+Z', 'Отменить последнее действие'),
    ('Redo', 'Ctrl+Y', 'Повторить отменённое действие'),
]
for i, (menu, key, desc) in enumerate(data):
    table.rows[i+1].cells[0].text = menu
    table.rows[i+1].cells[1].text = key
    table.rows[i+1].cells[2].text = desc

doc.add_heading('4.3. Grammar — операции с грамматикой', level=2)
doc.add_paragraph(
    'Это основное меню для работы с грамматикой. Все пункты также продублированы '
    'на панели операций справа.'
)
table = doc.add_table(rows=15, cols=2)
table.style = 'Table Grid'
for i, h in enumerate(['Пункт меню', 'Описание']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
data = [
    ('Parse (F5)', 'Разобрать текст грамматики и построить диаграммы'),
    ('Semantics...', 'Открыть окно управления семантическими действиями'),
    ('Eliminate Left Recursion', 'Устранить левую рекурсию для всех нетерминалов'),
    ('Eliminate Right Recursion', 'Устранить правую рекурсию для всех нетерминалов'),
    ('Regularize (Both)', 'Применить устранение и левой, и правой рекурсии'),
    ('Minimize (DFA)', 'Минимизация автомата (DFA) для всех нетерминалов'),
    ('Analyze Recursion', 'Открыть окно анализа рекурсий'),
    ('Extract Rule', 'Извлечь выделенный фрагмент диаграммы в новый нетерминал'),
    ('Substitute', 'Подставить определение нетерминала вместо ссылки'),
    ('Remove Useless', 'Удалить бесполезные (недостижимые/непродуктивные) символы'),
    ('Toggle Macro', 'Пометить/снять пометку нетерминала как макроса'),
    ('Open All Macros', 'Раскрыть все макроопределения в диаграммах'),
    ('Close All Definitions', 'Свернуть все раскрытые определения'),
    ('Macros Opened by Default', 'Переключить режим раскрытия макросов по умолчанию'),
]
for i, (menu, desc) in enumerate(data):
    table.rows[i+1].cells[0].text = menu
    table.rows[i+1].cells[1].text = desc

doc.add_heading('4.4. Help — справка', level=2)
doc.add_paragraph(
    'Пункт «Help» открывает встроенную справку с кратким описанием синтаксиса RBNF '
    'и списком горячих клавиш. Пункт «About» показывает информацию о версии приложения.'
)

doc.add_page_break()

# ── 5. Синтаксис RBNF ──
doc.add_heading('5. Синтаксис RBNF-грамматик', level=1)

doc.add_heading('5.1. Элементы грамматики', level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['Элемент', 'Синтаксис', 'Пример']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
data = [
    ('Терминал', "Одинарные кавычки: 'symbol'", "'if', '+', 'begin'"),
    ('Нетерминал', 'Идентификатор (буква + буквы/цифры)', 'stmt, expr, ID'),
    ('Макрос', 'Префикс @', '@list, @item'),
    ('Семантическое действие', 'Префикс $', '$action, $check'),
    ('Epsilon (пустое слово)', 'eps или @', 'S : eps.'),
]
for i, (el, syn, ex) in enumerate(data):
    table.rows[i+1].cells[0].text = el
    table.rows[i+1].cells[1].text = syn
    table.rows[i+1].cells[2].text = ex

doc.add_heading('5.2. Операторы', level=2)
table = doc.add_table(rows=8, cols=4)
table.style = 'Table Grid'
for i, h in enumerate(['Оператор', 'Символ', 'Описание', 'Пример']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
data = [
    ('Последовательность', ',', 'Элементы следуют друг за другом', "'a' , 'b' , 'c'"),
    ('Альтернатива', ';', 'Выбор одного из вариантов', "'a' ; 'b' ; 'c'"),
    ('Унарная итерация (*)', '*', 'Повторение 0 или более раз', "@*('item')"),
    ('Положит. итерация (+)', '+', 'Повторение 1 или более раз', "@+('digit')"),
    ('Бинарная итерация (#)', '#', 'Цикл с телом и разделителем', "item # ','"),
    ('Группировка', '( )', 'Приоритет операций', "('a' ; 'b') , 'c'"),
    ('Опциональная группировка', '[ ]', 'Опциональный блок (0 или 1 раз)', "['else' , stmt]"),
]
for i, (op, sym, desc, ex) in enumerate(data):
    table.rows[i+1].cells[0].text = op
    table.rows[i+1].cells[1].text = sym
    table.rows[i+1].cells[2].text = desc
    table.rows[i+1].cells[3].text = ex

doc.add_paragraph()
doc.add_paragraph(
    'Приоритет операторов (от высшего к низшему): '
    'группировка ( ) и [ ] → унарная итерация * + → бинарная итерация # → '
    'последовательность , → альтернатива ;'
)

doc.add_heading('5.3. Примеры грамматик', level=2)
doc.add_paragraph('Простая грамматика объявлений переменных:')
p = doc.add_paragraph()
run = p.add_run(
    "S : declarations.\n"
    "declarations : declaration # ','.\n"
    "declaration : type , identifier.\n"
    "type : 'int' ; 'float' ; 'string'.\n"
    "identifier : ID.\n"
    "EOGram!"
)
run.font.name = 'Courier New'
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph('Грамматика арифметических выражений (с левой рекурсией):')
p = doc.add_paragraph()
run = p.add_run(
    "expr : expr , '+' , term ; expr , '-' , term ; term.\n"
    "term : term , '*' , factor ; term , '/' , factor ; factor.\n"
    "factor : '(' , expr , ')' ; ID ; NUM.\n"
    "EOGram!"
)
run.font.name = 'Courier New'
run.font.size = Pt(11)

doc.add_page_break()

# ── 6. Работа с диаграммой ──
doc.add_heading('6. Работа с диаграммой', level=1)

doc.add_heading('6.1. Навигация и масштабирование', level=2)
bullets = [
    'Перетаскивание объекта: зажать левую кнопку мыши на элементе диаграммы и перетащить.',
    'Масштабирование: колёсико мыши увеличивает/уменьшает масштаб диаграммы.',
    'Прокрутка: зажать среднюю кнопку мыши (или колёсико) для панорамирования.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('6.2. Выделение объектов', level=2)
doc.add_paragraph(
    'Для выделения нескольких объектов: зажмите левую кнопку мыши на пустом месте '
    'диаграммы и протяните рамку выделения. Все объекты, попавшие в рамку, будут выделены. '
    'Выделенные объекты подсвечиваются. Выделение используется для операций Extract Rule и Delete.'
)

doc.add_heading('6.3. Контекстное меню диаграммы', level=2)
doc.add_paragraph(
    'Правый клик на элементе диаграммы или на пустом месте открывает контекстное меню:'
)
table = doc.add_table(rows=11, cols=2)
table.style = 'Table Grid'
for i, h in enumerate(['Пункт', 'Описание']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
data = [
    ('Edit Symbol', 'Редактировать имя терминала или нетерминала'),
    ('Delete Selected (Del)', 'Удалить выделенные объекты из диаграммы'),
    ('Extract Rule', 'Извлечь выделенный фрагмент в новый нетерминал'),
    ('Substitute', 'Заменить ссылку на нетерминал его определением'),
    ('Add (AND)', 'Добавить терминал в последовательность (через запятую)'),
    ('Add (OR)', 'Добавить терминал как альтернативу (через точку с запятой)'),
    ('Create New (NT)', 'Создать новый нетерминал и добавить ссылку'),
    ('Add Reference (AND)', 'Добавить ссылку на существующий нетерминал в последовательность'),
    ('Add Reference (OR)', 'Добавить ссылку на существующий нетерминал как альтернативу'),
    ('Add Extended Point', 'Добавить точку расширения диаграммы'),
]
for i, (menu, desc) in enumerate(data):
    table.rows[i+1].cells[0].text = menu
    table.rows[i+1].cells[1].text = desc

doc.add_page_break()

# ── 7. Преобразования ──
doc.add_heading('7. Преобразования грамматик', level=1)
doc.add_paragraph(
    'Все преобразования доступны через меню Grammar и панель операций справа.'
)

doc.add_heading('7.1. Устранение левой рекурсии (Eliminate Left Recursion)', level=2)
doc.add_paragraph(
    'Преобразует леворекурсивные правила вида A → Aα | β в эквивалентные '
    'нерекурсивные: A → βA\', A\' → αA\' | ε. Это необходимо для построения '
    'нисходящих синтаксических анализаторов (LL-парсеров).'
)

doc.add_heading('7.2. Устранение правой рекурсии (Eliminate Right Recursion)', level=2)
doc.add_paragraph(
    'Аналогично устранению левой рекурсии, но для праворекурсивных правил '
    'вида A → αA | β. Результат: A → α*β — используется итерация.'
)

doc.add_heading('7.3. Регуляризация (Regularize)', level=2)
doc.add_paragraph(
    'Применяет одновременно устранение и левой, и правой рекурсии для всех нетерминалов.'
)

doc.add_heading('7.4. Минимизация (Minimize DFA)', level=2)
doc.add_paragraph(
    'Строит детерминированный конечный автомат (DFA) для каждого правила и '
    'минимизирует его. Результат — эквивалентное, но более компактное регулярное выражение.'
)

doc.add_heading('7.5. Извлечение правила (Extract Rule)', level=2)
doc.add_paragraph(
    'Позволяет выделить фрагмент диаграммы (один или несколько связанных объектов) '
    'и вынести его в новый нетерминал. Порядок действий:'
)
steps = [
    'Выделите один или несколько объектов на диаграмме рамкой выделения.',
    'Выберите Grammar → Extract Rule (или правый клик → Extract Rule).',
    'В появившемся диалоге задайте имя нового нетерминала (предлагается автоматически).',
    'Нажмите OK — выделенный фрагмент будет заменён ссылкой на новый нетерминал.',
]
for i, s in enumerate(steps):
    doc.add_paragraph(f'{i+1}. {s}')

doc.add_heading('7.6. Подстановка (Substitute)', level=2)
doc.add_paragraph(
    'Обратная операция к Extract Rule: выделите ссылку на нетерминал и выберите Substitute. '
    'Ссылка будет заменена определением (телом правила) этого нетерминала.'
)

doc.add_heading('7.7. Удаление бесполезных символов (Remove Useless)', level=2)
doc.add_paragraph(
    'Удаляет из грамматики нетерминалы, которые недостижимы из начального символа S, '
    'а также непродуктивные символы (не порождающие ни одной терминальной цепочки).'
)

doc.add_heading('7.8. Левая факторизация (Left Factorization)', level=2)
doc.add_paragraph(
    'Выделяет общий префикс из альтернатив: A → αβ | αγ преобразуется в A → αA\', '
    'A\' → β | γ. Доступна через консольное приложение syngt_cli.'
)

doc.add_page_break()

# ── 8. Анализ ──
doc.add_heading('8. Анализ грамматик', level=1)

doc.add_heading('8.1. Анализ рекурсий (Analyze Recursion)', level=2)
doc.add_paragraph(
    'Открывает окно анализа рекурсий, в котором для каждого нетерминала показано:'
)
bullets = [
    'Left Recursion — наличие левой рекурсии (direct/indirect/отсутствует).',
    'Recursion — наличие рекурсии в целом (любая позиция).',
    'Right Recursion — наличие правой рекурсии.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')
doc.add_paragraph(
    'В окне анализа доступны кнопки «Elim All Left» и «Elim All Right» '
    'для быстрого устранения рекурсии прямо из этого окна.'
)

doc.add_heading('8.2. Проверка LL(1) (Check LL(1))', level=2)
doc.add_paragraph(
    'Проверяет, является ли грамматика LL(1)-грамматикой. Результат выводится в консоль. '
    'Если грамматика не LL(1), указываются конфликты. Доступна через syngt_cli (check-ll1).'
)

doc.add_heading('8.3. Множества FIRST и FOLLOW', level=2)
doc.add_paragraph(
    'Вычисляет множества FIRST и FOLLOW для каждого нетерминала. '
    'Результаты выводятся в консоль. Доступно через syngt_cli (first-follow).'
)

# ── 9. Вспомогательные функции ──
doc.add_heading('9. Вспомогательные функции', level=1)

doc.add_heading('9.1. Макросы', level=2)
doc.add_paragraph(
    'Нетерминал можно пометить как макрос (Grammar → Toggle Macro). Макросы '
    'отображаются пунктирным прямоугольником на диаграмме. Их определение можно '
    'раскрыть (Open All Macros) или свернуть (Close All Definitions) для '
    'компактности диаграммы.'
)

doc.add_heading('9.2. Семантические действия', level=2)
doc.add_paragraph(
    'Окно семантик (Grammar → Semantics...) позволяет управлять семантическими '
    'действиями грамматики. Можно добавлять новые действия и удалять существующие. '
    'Семантические действия отображаются на стрелках диаграммы.'
)

doc.add_heading('9.3. Система Undo/Redo', level=2)
doc.add_paragraph(
    'Все изменения грамматики (редактирование правил, применение преобразований, '
    'добавление/удаление элементов) записываются в историю. Отмена: Ctrl+Z, повтор: Ctrl+Y.'
)

doc.add_page_break()

# ── 10. Горячие клавиши ──
doc.add_heading('10. Горячие клавиши', level=1)
table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
for i, h in enumerate(['Комбинация', 'Действие']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
data = [
    ('Ctrl+N', 'Новая грамматика'),
    ('Ctrl+O', 'Открыть файл'),
    ('Ctrl+S', 'Сохранить файл'),
    ('Ctrl+Z', 'Отмена'),
    ('Ctrl+Y', 'Повтор'),
    ('F5', 'Разобрать грамматику (Parse)'),
    ('Del', 'Удалить выделенные объекты'),
]
for i, (key, action) in enumerate(data):
    table.rows[i+1].cells[0].text = key
    table.rows[i+1].cells[1].text = action

doc.add_page_break()

# ── 11. CLI ──
doc.add_heading('11. Консольное приложение syngt_cli', level=1)
doc.add_paragraph(
    'Консольное приложение позволяет выполнять операции с грамматиками из командной строки '
    'без графического интерфейса. Формат вызова:'
)
p = doc.add_paragraph()
run = p.add_run('    syngt_cli <команда> [аргументы]')
run.font.name = 'Courier New'
run.font.size = Pt(11)
doc.add_paragraph()

table = doc.add_table(rows=9, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['Команда', 'Аргументы', 'Описание']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
data = [
    ('info', '<grammar.grm>', 'Вывести структуру грамматики: терминалы, нетерминалы, правила'),
    ('regularize', '<in.grm> <out.grm>', 'Полная регуляризация (устранение рекурсий + факторизация + удаление бесполезных)'),
    ('eliminate-left', '<in.grm> <out.grm>', 'Устранить только левую рекурсию'),
    ('factorize', '<in.grm> <out.grm>', 'Левая факторизация'),
    ('remove-useless', '<in.grm> <out.grm>', 'Удалить бесполезные символы'),
    ('check-ll1', '<grammar.grm>', 'Проверить LL(1)-свойство'),
    ('first-follow', '<grammar.grm>', 'Вычислить множества FIRST и FOLLOW'),
    ('table', '<grammar.grm>', 'Построить таблицу разбора LL(1)'),
]
for i, (cmd, args, desc) in enumerate(data):
    table.rows[i+1].cells[0].text = cmd
    table.rows[i+1].cells[1].text = args
    table.rows[i+1].cells[2].text = desc

doc.add_paragraph()
doc.add_paragraph('Пример использования:')
p = doc.add_paragraph()
run = p.add_run(
    '    syngt_cli info grammar.grm\n'
    '    syngt_cli regularize input.grm output.grm\n'
    '    syngt_cli check-ll1 grammar.grm\n'
    '    syngt_cli first-follow grammar.grm'
)
run.font.name = 'Courier New'
run.font.size = Pt(11)

# Save
doc.save('/mnt/d/geXrBy/Documents/vkr/Руководство_пользователя_SynGT.docx')
print('User guide saved.')
