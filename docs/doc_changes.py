#!/usr/bin/env python3
"""Generate 'What's New + Tests' DOCX for SynGT C++ application."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(18)
    elif level == 2:
        hs.font.size = Pt(16)
    else:
        hs.font.size = Pt(14)

# ── Title ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Описание изменений\nв системе SynGT (C++) относительно исходной версии')
run.bold = True
run.font.size = Pt(22)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Новый функционал и покрытие тестами')
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

doc.add_paragraph()

# ── TOC ──
doc.add_heading('Содержание', level=1)
toc = [
    'Часть I. Изменения относительно предыдущей версии',
    '  1. Архитектурные изменения',
    '  2. Новый функционал',
    '  3. Исправленные ошибки',
    '  4. Улучшения интерфейса',
    'Часть II. Тестирование',
    '  5. Структура тестов',
    '  6. Покрытие по модулям',
    '  7. Детальное описание тестов',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PART I
# ══════════════════════════════════════════════════════════

doc.add_heading('Часть I. Изменения относительно предыдущей версии', level=1)
doc.add_paragraph(
    'Ниже описаны все значимые изменения, внесённые в систему SynGT (C++) '
    'по сравнению с версией, описанной в отчёте по учебной практике (ноябрь 2025 г.).'
)

# ── 1 ──
doc.add_heading('1. Архитектурные изменения', level=1)

doc.add_heading('1.1. Устранение правой рекурсии (RightElimination)', level=2)
doc.add_paragraph(
    'Добавлен модуль устранения правой рекурсии (RightElimination), '
    'симметричный модулю устранения левой рекурсии. Исходная версия '
    'поддерживала только устранение левой рекурсии. Теперь доступна полная '
    'регуляризация: устранение и левой, и правой рекурсии за один вызов (Regularize).'
)

doc.add_heading('1.2. Анализ рекурсий (RecursionAnalyzer)', level=2)
doc.add_paragraph(
    'Добавлен модуль анализа рекурсий, портированный из оригинальной Pascal-версии '
    '(модуль Analyzer.pas). Алгоритм использует BFS по графу вызовов нетерминалов '
    'и определяет для каждого нетерминала тип рекурсии:'
)
bullets = [
    'Левая рекурсия (direct/indirect) — нетерминал может вызвать себя в крайней левой позиции.',
    'Правая рекурсия (direct/indirect) — нетерминал может вызвать себя в крайней правой позиции.',
    'Рекурсия в целом (direct/indirect) — нетерминал может вызвать себя в любой позиции.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'В ходе портирования были выявлены и исправлены расхождения с оригинальным алгоритмом:'
)
bullets = [
    'Исправлена обработка семантических действий в canProduceEpsilon — '
    'семантики теперь корректно трактуются как прозрачные (epsilon-порождающие), '
    'поскольку не потребляют входных символов.',
    'Исправлен разбор ключевого слова eps — ранее eps ошибочно добавлялся как нетерминал, '
    'теперь корректно распознаётся как терминал-epsilon (RETerminal с ID=0).',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('1.3. Извлечение правила (Extract Rule)', level=2)
doc.add_paragraph(
    'Реализована функция извлечения фрагмента диаграммы в новый нетерминал. '
    'Алгоритм находит минимальное поддерево RE-дерева, покрывающее все выделенные объекты '
    '(FindSelectionCover), и заменяет его ссылкой на новый нетерминал.'
)
doc.add_paragraph(
    'Добавлена логика расширения покрытия (ExpandSemanticsCover): если выделен один '
    'элемент (терминал или нетерминал), то соседние семантические действия в цепочке AND '
    'автоматически включаются в извлекаемый фрагмент. Это обеспечивает корректную '
    'работу с грамматиками, содержащими семантические действия.'
)

doc.add_heading('1.4. Подстановка (Substitute)', level=2)
doc.add_paragraph(
    'Реализована операция подстановки — замена ссылки на нетерминал его определением. '
    'Обратная операция к Extract Rule.'
)

doc.add_heading('1.5. Окно семантик', level=2)
doc.add_paragraph(
    'Семантические действия вынесены из основной панели в отдельное плавающее окно '
    '(Grammar → Semantics...). Окно позволяет добавлять и удалять семантические действия.'
)

doc.add_page_break()

# ── 2 ──
doc.add_heading('2. Новый функционал', level=1)

doc.add_heading('2.1. Пустая грамматика при запуске', level=2)
doc.add_paragraph(
    'При запуске приложения и при создании новой грамматики (File → New) '
    'вместо пустого экрана теперь создаётся начальная грамматика «S : eps.» '
    'с визуализацией в виде стрелки от начала к концу (epsilon-правило). '
    'Ранее программа запускалась с пустым состоянием без визуализации.'
)

doc.add_heading('2.2. Визуализация epsilon-правил', level=2)
doc.add_paragraph(
    'Epsilon-правила (правила вида S : eps.) теперь корректно визуализируются '
    'в виде диаграммы со стрелкой от начального треугольника к конечному, '
    'вместо пустой области.'
)

doc.add_heading('2.3. Исправление парсинга ключевого слова eps', level=2)
doc.add_paragraph(
    'Парсер теперь корректно распознаёт ключевое слово «eps» как пустое слово '
    '(epsilon, RETerminal с ID=0). Ранее eps ошибочно обрабатывался как '
    'идентификатор нетерминала, что приводило к созданию фиктивного нетерминала «eps» '
    'и некорректным результатам анализа.'
)

doc.add_heading('2.4. Окно анализа рекурсий', level=2)
doc.add_paragraph(
    'Добавлено интерактивное окно анализа рекурсий (Grammar → Analyze Recursion), '
    'отображающее таблицу с типами рекурсий для каждого нетерминала. '
    'Окно содержит кнопки для быстрого устранения левой и правой рекурсии, '
    'а также кнопку «Refresh» для обновления анализа после изменений.'
)

doc.add_heading('2.5. Устранение правой рекурсии в GUI', level=2)
doc.add_paragraph(
    'В меню Grammar и на панель операций добавлены пункты:'
)
bullets = [
    'Eliminate Right Recursion — устранение только правой рекурсии.',
    'Regularize (Both) — устранение обоих видов рекурсии за один вызов.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('2.6. Контекстное меню диаграммы', level=2)
doc.add_paragraph(
    'Расширено контекстное меню (правый клик на диаграмме). Добавлены пункты: '
    'Extract Rule, Substitute, Add Extended Point. Ранее контекстное меню содержало '
    'только базовые операции редактирования (Edit Symbol, Remove Symbol, Add Terminal/NT).'
)

doc.add_page_break()

# ── 3 ──
doc.add_heading('3. Исправленные ошибки', level=1)

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['Проблема', 'Где', 'Исправление']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
data = [
    ('eps парсился как нетерминал',
     'Parser.cpp, parseK()',
     'Добавлена проверка if (name == "eps") перед блоком общих идентификаторов; возвращается RETerminal(grammar, 0)'),
    ('Семантики мешали анализу рекурсий',
     'RecursionAnalyzer.cpp, canProduceEpsilon()',
     'RESemantic теперь возвращает true (прозрачна для epsilon), а не false'),
    ('Extract Rule захватывал всю строку при выделении 2+ объектов',
     'main.cpp, ExpandSemanticsCover()',
     'Расширение покрытия теперь применяется только для одиночных листовых узлов (RETerminal/RENonTerminal)'),
    ('buildRestoredText пропускал epsilon-правила',
     'main.cpp, buildRestoredText()',
     'Для пустых правил генерируется "name : eps.\\n" вместо пропуска'),
]
for i, (prob, where, fix) in enumerate(data):
    table.rows[i+1].cells[0].text = prob
    table.rows[i+1].cells[1].text = where
    table.rows[i+1].cells[2].text = fix

# ── 4 ──
doc.add_heading('4. Улучшения интерфейса', level=1)
bullets = [
    'Семантики вынесены из правой панели в отдельное плавающее окно, освобождая место для панели операций.',
    'При запуске отображается начальная диаграмма (S : eps) вместо пустого экрана.',
    'Добавлено окно анализа рекурсий с кнопками быстрого устранения.',
    'Расширено контекстное меню диаграммы (Extract Rule, Substitute).',
    'ParseGrammar упрощён: activeNTIndex больше не пропускает epsilon-нетерминалы.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PART II
# ══════════════════════════════════════════════════════════

doc.add_heading('Часть II. Тестирование', level=1)

# ── 5 ──
doc.add_heading('5. Структура тестов', level=1)
doc.add_paragraph(
    'Система тестирования построена на базе Google Test. Тесты организованы '
    'по категориям, соответствующим модулям библиотеки libsyngt. '
    'Каждый тестовый файл (test_*.cpp) компилируется в отдельный исполняемый модуль.'
)
doc.add_paragraph(
    'Общая статистика:'
)
table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
for i, h in enumerate(['Метрика', 'Значение']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
data = [
    ('Тестовых файлов', '25'),
    ('Индивидуальных тестов', '~180'),
    ('Категорий', '7 (analysis, core, graphics, integration, parser, regex, transform, utils)'),
    ('Фреймворк', 'Google Test 1.14.0'),
    ('Платформы', 'Windows (MSVC, MinGW), Linux (GCC, Clang), macOS (Apple Clang)'),
    ('Время выполнения', '< 2 секунд'),
    ('Покрытие модулей', '100% модулей библиотеки libsyngt'),
]
for i, (m, v) in enumerate(data):
    table.rows[i+1].cells[0].text = m
    table.rows[i+1].cells[1].text = v

# ── 6 ──
doc.add_heading('6. Покрытие по модулям', level=1)

doc.add_paragraph(
    'В таблице ниже показано соответствие между модулями библиотеки, '
    'тестируемыми функциями и тестовыми файлами.'
)

# Big coverage table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
for i, h in enumerate(['Модуль', 'Тестируемые классы/функции', 'Тестовый файл', 'Кол-во тестов']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ('Core', 'Grammar: создание, загрузка, сохранение, управление списками', 'test_Grammar.cpp', '7'),
    ('Core', 'TerminalList, NonTerminalList, SemanticList, MacroList', 'test_Lists.cpp', '4'),
    ('Core', 'NTListItem: разбор правил, RE-дерево, copy, макросы', 'test_NTListItem.cpp', '8'),
    ('Core', 'NTListItem + Grammar: макросы, флаги, циклические ссылки', 'test_Macros.cpp', '9'),
    ('Parser', 'CharProducer: символьный поток, навигация, сброс', 'test_CharProducer.cpp', '8'),
    ('Parser', 'Parser: разбор последовательностей, альтернатив, итераций, группировок', 'test_Parser.cpp', '5'),
    ('Regex', 'RETree: создание узлов, копирование, toString', 'test_RETree.cpp', '11'),
    ('Regex', 'REAnd, REOr, REIteration: вложенные конструкции, семантики', 'test_REOperations.cpp', '7'),
    ('Transform', 'LeftElimination: обнаружение и устранение левой рекурсии', 'test_LeftElimination.cpp', '11'),
    ('Transform', 'RightElimination: обнаружение и устранение правой рекурсии', 'test_RightElimination.cpp', '13'),
    ('Transform', 'LeftFactorization: выделение общих префиксов', 'test_LeftFactorization.cpp', '10'),
    ('Transform', 'RemoveUseless: удаление недостижимых/непродуктивных', 'test_GrammarTransformations.cpp', '10'),
    ('Transform', 'FirstFollow: FIRST/FOLLOW, проверка LL(1)', 'test_FirstFollow.cpp', '15'),
    ('Transform', 'Regularize: полная регуляризация', 'test_Regularize.cpp', '7'),
    ('Analysis', 'RecursionAnalyzer: анализ рекурсий (left/right/any)', 'test_RecursionAnalyzer.cpp', '7'),
    ('Analysis', 'ParsingTable: построение таблиц LL(1), конфликты', 'test_ParsingTable.cpp', '7'),
    ('Analysis', 'Minimize: DFA-минимизация', 'test_Minimize.cpp', '7'),
    ('Graphics', 'Arrow: конструкторы, копирование, сериализация, семантики', 'test_Arrow.cpp', '13'),
    ('Graphics', 'DrawPoint: координаты, перемещение, виртуальные методы', 'test_DrawPoint.cpp', '12'),
    ('Utils', 'Creator: создание диаграмм из RE-деревьев', 'test_Creator.cpp', '6'),
    ('Utils', 'SemanticIDList: добавление, копирование, сериализация', 'test_Semantic.cpp', '8'),
    ('Utils', 'UndoRedo: отмена, повтор, история, дедупликация', 'test_UndoRedo.cpp', '17'),
    ('Integration', 'Grammar persistence: полный цикл сохранения/загрузки', 'test_GrammarPersistence.cpp', '1'),
    ('Integration', 'Загрузка реального файла LANG.GRM', 'test_LoadGrammar.cpp', '1'),
    ('Integration', 'Parse → Transform → Visualize: сквозные сценарии', 'test_Integration.cpp', '8'),
]
for mod, func, test_file, count in data:
    row = table.add_row()
    row.cells[0].text = mod
    row.cells[1].text = func
    row.cells[2].text = test_file
    row.cells[3].text = count

doc.add_page_break()

# ── 7. Detailed tests ──
doc.add_heading('7. Детальное описание тестов', level=1)

# Analysis
doc.add_heading('7.1. Категория analysis', level=2)

doc.add_heading('test_RecursionAnalyzer.cpp', level=3)
doc.add_paragraph('Тесты анализатора рекурсий (RecursionAnalyzer::analyze):')
tests = [
    ('NoRecursion', 'Грамматика без рекурсий: проверяет, что все поля left/right/any пусты.'),
    ('DirectLeftRecursion', 'Грамматика E → E \'+\' term ; term: обнаружение прямой левой рекурсии.'),
    ('DirectRightRecursion', 'Грамматика A → \'a\' A ; \'b\': обнаружение прямой правой рекурсии.'),
    ('MixedGrammar', 'Грамматика с рекурсивными и нерекурсивными нетерминалами одновременно.'),
    ('ResultCountMatchesNTs', 'Проверяет, что результатов ровно столько, сколько нетерминалов.'),
    ('EmptyGrammarDoesNotCrash', 'Пустая грамматика (nullptr) не вызывает падение.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_ParsingTable.cpp', level=3)
doc.add_paragraph('Тесты построения таблиц разбора LL(1):')
tests = [
    ('BuildSimpleTable', 'Построение таблицы для простого правила с одним терминалом.'),
    ('BuildWithAlternatives', 'Построение таблицы для правил с несколькими альтернативами.'),
    ('NoConflictsLL1Grammar', 'Проверка, что LL(1)-грамматика не порождает конфликтов.'),
    ('DetectConflicts', 'Обнаружение FIRST/FIRST конфликтов в не-LL(1) грамматике.'),
    ('PrintDoesNotCrash', 'Вывод таблицы не вызывает падение.'),
    ('ExportForCodegen', 'Экспорт таблицы в формат для генерации кода.'),
    ('ComplexGrammar', 'Таблица для реальной грамматики выражений (E, E\', T).'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_Minimize.cpp', level=3)
doc.add_paragraph('Тесты DFA-минимизации:')
tests = [
    ('SingleTerminalDoesNotCrash', 'Минимизация одиночного терминала не вызывает падение.'),
    ('AlternativeTerminals', 'Минимизация альтернативы \'a\' ; \'b\' ; \'c\'.'),
    ('Sequence', 'Минимизация последовательности (цепочки конкатенаций).'),
    ('AllNTsPreserved', 'Все нетерминалы сохраняются после минимизации.'),
    ('EmptyGrammarDoesNotCrash', 'Пустая грамматика обрабатывается без ошибок.'),
    ('RegularizedGrammarCanBeMinimized', 'Грамматика после регуляризации успешно минимизируется.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_page_break()

# Core
doc.add_heading('7.2. Категория core', level=2)

doc.add_heading('test_Grammar.cpp', level=3)
doc.add_paragraph('Тесты основного класса Grammar:')
tests = [
    ('Initialization', 'Создание пустой грамматики через fillNew(): начальный нетерминал S, списки пусты.'),
    ('AddTerminals', 'Добавление терминалов: проверка уникальности и ID.'),
    ('FindTerminals', 'Поиск терминалов по имени: найденные и ненайденные.'),
    ('AddMultipleTypes', 'Одновременное добавление терминалов, нетерминалов, семантик.'),
    ('SaveAndLoad', 'Полный цикл: создание → заполнение → сохранение в файл → загрузка → проверка.'),
    ('LoadRegularizeSave', 'Загрузка LANG.GRM → регуляризация → сохранение → проверка целостности.'),
    ('LoadLANGGRM', 'Загрузка файла с 50+ терминалами и нетерминалами.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_Lists.cpp', level=3)
doc.add_paragraph('Тесты списков символов:')
tests = [
    ('TerminalList::AddAndFind', 'Добавление и поиск терминалов, проверка дедупликации.'),
    ('NonTerminalList::FillNewAndAdd', 'Создание списка с нетерминалом S по умолчанию.'),
    ('SemanticList::BasicOperations', 'Добавление и получение семантических действий.'),
    ('MacroList::BasicOperations', 'Работа с макроопределениями.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_NTListItem.cpp', level=3)
doc.add_paragraph('Тесты элемента списка нетерминалов:')
tests = [
    ('SetValueParsesToTree', 'Разбор строки правила в RE-дерево.'),
    ('SetRootUpdatesValue', 'Обновление строки правила из RE-дерева (обратное преобразование).'),
    ('CopyRETree', 'Глубокое копирование RE-дерева.'),
    ('MacroFlagDefault', 'Проверка, что по умолчанию NT не является макросом.'),
    ('SetMacroTrue / SetMacroFalse', 'Установка и снятие флага макроса.'),
    ('SetMarkInProgress', 'Переходы состояний пометки (mark).'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_Macros.cpp', level=3)
doc.add_paragraph('Тесты системы макросов:')
tests = [
    ('OpenMacroRefsDefaultTrue/False', 'Раскрытие/свёртывание макросов в диаграмме.'),
    ('DrawObjectTypeForMacro', 'Макрос отображается как ctDrawObjectMacro.'),
    ('DrawObjectTypeForNonMacro', 'Обычный NT отображается как ctDrawObjectNonTerminal.'),
    ('AllMacroWasOpenedNormal', 'Проверка, что все макросы раскрыты.'),
    ('AllMacroWasOpenedCircularThrows', 'Обнаружение циклических ссылок между макросами (A→B, B→A).'),
    ('SaveLoadPreservesMacroFlag', 'Сохранение и загрузка флага макроса.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_page_break()

# Parser
doc.add_heading('7.3. Категория parser', level=2)

doc.add_heading('test_CharProducer.cpp', level=3)
doc.add_paragraph('Тесты символьного потока (CharProducer):')
tests = [
    ('InitialChar', 'Начальный символ и позиция (index=0).'),
    ('AdvanceNext', 'Последовательный проход по символам.'),
    ('EndDetection', 'Обнаружение конца строки.'),
    ('NextReturnsFalseAtEnd', 'next() возвращает false при достижении конца.'),
    ('EmptyString', 'Обработка пустой строки.'),
    ('Reset', 'Сброс позиции к началу.'),
    ('GetString', 'Доступ к полной строке.'),
    ('TraverseFullString', 'Полный проход по строке до конца.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_Parser.cpp', level=3)
doc.add_paragraph('Тесты парсера RBNF:')
tests = [
    ('SimpleSequence', 'Разбор последовательности: \'begin\' , \'end\'.'),
    ('SimpleAlternative', 'Разбор альтернативы: \'program\' ; \'begin\'.'),
    ('Iteration', 'Разбор итерации: statement * expression.'),
    ('WithParentheses', 'Разбор группировки со скобками.'),
    ('ComplexExpression', 'Разбор сложных вложенных конструкций.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

# Transform
doc.add_heading('7.4. Категория transform', level=2)

doc.add_heading('test_LeftElimination.cpp', level=3)
doc.add_paragraph('Тесты устранения левой рекурсии (11 тестов):')
tests = [
    ('DetectDirectLeftRecursion', 'Обнаружение E → E \'+\' term ; term.'),
    ('NoLeftRecursionInRightRecursive', 'Правая рекурсия не ошибочно определяется как левая.'),
    ('NoLeftRecursionInPlainRule', 'Нерекурсивные правила не затрагиваются.'),
    ('EliminateSimpleLeftRecursion', 'E → E \'+\' term ; term → term(\'+\' term)*.'),
    ('EliminateMultipleAlternatives', 'E → E \'+\' ; E \'-\' ; num — несколько рекурсивных альтернатив.'),
    ('EliminateForWholeGrammar', 'Устранение для всей грамматики целиком.'),
    ('PureLeftRecursion_NoBase', 'Только рекурсивные альтернативы, нет базового случая.'),
    ('EpsilonInAlternative', 'S → S \'a\' | @ — альтернатива с epsilon.'),
    ('SaveAndLoadAfterElimination', 'Сохранение и загрузка после преобразования.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_RightElimination.cpp', level=3)
doc.add_paragraph('Тесты устранения правой рекурсии (13 тестов):')
tests = [
    ('DetectDirectRightRecursion', 'Обнаружение A → \'a\' A ; \'b\'.'),
    ('NoRightRecursionInLeftRecursive', 'Левая рекурсия не ошибочно определяется как правая.'),
    ('EliminateSimpleRightRecursion', 'A → \'a\' A ; \'b\' → \'b\' \'a\'*.'),
    ('EliminateMultipleAlternatives', 'Несколько праворекурсивных альтернатив.'),
    ('PureRightRecursion_NoBase', 'Только рекурсивные альтернативы.'),
    ('EpsilonInAlternative', 'A → \'a\' A | @ — альтернатива с epsilon.'),
    ('SymmetryWithLeftElimination', 'Сравнение с устранением левой рекурсии: List → \'item\' List | @.'),
    ('SaveAndLoadAfterElimination', 'Сохранение и загрузка после преобразования.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_LeftFactorization.cpp', level=3)
doc.add_paragraph('Тесты левой факторизации (10 тестов):')
tests = [
    ('DetectCommonPrefix', 'Обнаружение общего префикса \'if\' expr \'then\' ...'),
    ('NoCommonPrefix', '\'a\' | \'b\' — нет общего префикса, пропуск.'),
    ('FactorizeSimple', '\'a\' \'b\' | \'a\' \'c\' → \'a\' A\', A\' → \'b\' | \'c\'.'),
    ('FactorizeMultiple', '\'x\' \'a\' | \'x\' \'b\' | \'x\' \'c\' — три альтернативы с общим префиксом.'),
    ('FactorizeAllGrammar', 'Факторизация для всех NT одной грамматики.'),
    ('RecursiveFactorization', 'Вложенные общие префиксы.'),
    ('DeepRecursion', 'Многоуровневая вложенность.'),
    ('NoInfiniteLoop', 'Гарантия завершения алгоритма.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_FirstFollow.cpp', level=3)
doc.add_paragraph('Тесты FIRST/FOLLOW множеств и проверки LL(1) (15 тестов):')
tests = [
    ('FirstSingleTerminal', 'FIRST({S}) для S → \'a\'.'),
    ('FirstSequence', 'FIRST для последовательности \'a\',\'b\'.'),
    ('FirstAlternatives', 'Объединение FIRST по альтернативам.'),
    ('FirstWithNonTerminal', 'Транзитивное вычисление через нетерминал.'),
    ('FirstRecursiveGrammar', 'E → T E\', E\' → \'+\' T E\' | @ — с epsilon-продукциями.'),
    ('FollowStartSymbol', 'FOLLOW(S) содержит маркер конца.'),
    ('FollowBasic', 'FOLLOW(A) = {\'b\'} для S → A \'b\'.'),
    ('LL1Valid', 'Корректная LL(1)-грамматика проходит проверку.'),
    ('LL1Invalid_FirstConflict', 'Обнаружение конфликта (\'a\' \'b\' | \'a\' \'c\').'),
    ('LL1AfterFactorization', 'После факторизации грамматика становится LL(1).'),
    ('ComplexGrammar', 'Грамматика с epsilon-продукциями.'),
    ('PrintSetsDoesNotCrash', 'Вывод множеств не вызывает падение.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_Regularize.cpp', level=3)
doc.add_paragraph('Тесты полной регуляризации (7 тестов):')
tests = [
    ('SimpleLeftRecursiveGrammar', 'Регуляризация грамматики с левой рекурсией.'),
    ('SimpleRightRecursiveGrammar', 'Регуляризация грамматики с правой рекурсией.'),
    ('PlainGrammarUnchanged', 'Нерекурсивная грамматика не изменяется.'),
    ('ClassicExpressionGrammar', 'Полная грамматика выражений (expr/term/factor).'),
    ('EmptyGrammarDoesNotCrash', 'Пустая грамматика обрабатывается без ошибок.'),
    ('MultipleNTs_AllPreserved', 'Все нетерминалы сохраняются после регуляризации.'),
    ('SameResultAsGrammarMethod', 'Результат совпадает с Grammar::regularize().'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_page_break()

doc.add_heading('test_GrammarTransformations.cpp', level=3)
doc.add_paragraph('Интеграционные тесты преобразований (10 тестов):')
tests = [
    ('RegularizeSimpleGrammar', 'Регуляризация классической грамматики E, T, F.'),
    ('ChainedTransformations', 'Цепочка: факторизация → удаление бесполезных.'),
    ('IndirectLeftRecursion', 'S→A\'a\';\'b\', A→S\'c\';\'d\' — косвенная рекурсия.'),
    ('MultipleLeftFactors', 'Множественные общие префиксы.'),
    ('UnreachableNonTerminals', 'Удаление недостижимых нетерминалов.'),
    ('NonProductiveNonTerminals', 'Удаление непродуктивных циклических цепочек.'),
    ('LL1Check', 'Проверка LL(1)-свойства корректной грамматики.'),
    ('NotLL1Check', 'Обнаружение не-LL(1) грамматики.'),
    ('MakeLL1ByFactorization', 'Превращение не-LL(1) в LL(1) через факторизацию.'),
    ('ComplexRealWorldGrammar', 'Полная грамматика выражений из реального кода.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

# Graphics
doc.add_heading('7.5. Категория graphics', level=2)

doc.add_heading('test_Arrow.cpp', level=3)
doc.add_paragraph('Тесты класса Arrow — стрелки диаграммы (13 тестов):')
tests = [
    ('DefaultConstructor', 'Создание стрелки с длиной MinArrowLength.'),
    ('ConstructorWithParameters', 'Стрелка с направлением (ward) и источником.'),
    ('ConstructorWithNoneWard', 'Стрелка без направления.'),
    ('SetFromDO / SetWard', 'Установка свойств.'),
    ('Copy', 'Глубокое копирование с сохранением ward и source.'),
    ('Save', 'Сериализация (ctArrow=0, код направления).'),
    ('SemanticArrowConstructor', 'Стрелка с семантическими действиями.'),
    ('GetSemantics', 'Получение списка семантических ID.'),
    ('SemanticArrowCopy / SemanticArrowSave', 'Копирование и сериализация с семантиками.'),
    ('NullSemantics', 'Обработка null-семантик.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_DrawPoint.cpp', level=3)
doc.add_paragraph('Тесты класса DrawPoint — базовый позиционируемый объект (12 тестов):')
tests = [
    ('InitialCoordinates', 'Начальные координаты (0,0), endX смещено на длину.'),
    ('SetPosition', 'Установка произвольной позиции.'),
    ('SetXCoord / SetYCoord', 'Установка отдельных осей.'),
    ('Move / MoveNegative', 'Относительное перемещение (положительное и отрицательное).'),
    ('DifferentLengths', 'Объекты с разной длиной.'),
    ('VirtualSetters', 'Виртуальные методы правильно вызываются.'),
    ('ChainedOperations', 'Цепочка операций перемещения.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

# Regex
doc.add_heading('7.6. Категория regex', level=2)

doc.add_heading('test_RETree.cpp', level=3)
doc.add_paragraph('Тесты базового класса RE-дерева (11 тестов):')
tests = [
    ('TerminalCreation / TerminalCopy', 'Создание и копирование RETerminal.'),
    ('AndOperation', 'Создание REAnd (последовательность через \',\').'),
    ('OrOperation', 'Создание REOr (альтернатива через \';\').'),
    ('IterationOperation', 'Создание REIteration (звёздочка Клини \'*\').'),
    ('ComplexExpression', 'Вложенные операции с нетерминалами.'),
    ('Semantics', 'Создание узлов семантических действий.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_REOperations.cpp', level=3)
doc.add_paragraph('Тесты составных операций RE-дерева (7 тестов):')
tests = [
    ('DeepNesting', 'Глубоко вложенные конструкции ((\'a\',\'b\');\'c\'),\'d\'.'),
    ('MultipleIterations', 'Вложенные итерации \'a\' * (\'b\' * \'c\').'),
    ('MixedNonTerminalsAndTerminals', 'A , \'plus\' , B.'),
    ('SemanticsInExpression', '\'a\' , @action , \'b\'.'),
    ('EmptyAlternative', '\'a\' | @.'),
    ('CopyPreservesStructure', 'Глубокое копирование сохраняет структуру.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

# Utils
doc.add_heading('7.7. Категория utils', level=2)

doc.add_heading('test_Creator.cpp', level=3)
doc.add_paragraph('Тесты создания диаграмм из RE-деревьев (6 тестов):')
tests = [
    ('CreateEmptyDiagram', 'Визуализация одиночного терминала.'),
    ('CreateSequence', 'Горизонтальное размещение последовательности.'),
    ('CreateAlternative', 'Размещение альтернативы с точками разветвления.'),
    ('DiagramSizePositive', 'Диаграмма имеет положительные размеры (width, height > 0).'),
    ('MultipleCreations', 'Повторное создание диаграммы идемпотентно.'),
    ('NullPointerHandling', 'Обработка null-параметров без падения.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_Semantic.cpp', level=3)
doc.add_paragraph('Тесты списка семантических ID (8 тестов):')
tests = [
    ('DefaultConstruction', 'Пустой список при создании.'),
    ('AddElements', 'Добавление и получение ID по индексу.'),
    ('Clear', 'Очистка списка.'),
    ('GetLength', 'Вычисление визуальной длины (50 на семантику).'),
    ('Copy', 'Глубокое копирование с независимостью.'),
    ('GetItems', 'Доступ к внутреннему вектору.'),
    ('SaveLoad', 'Сериализация и десериализация.'),
    ('MultipleAdds', 'Стресс-тест с 10 семантиками.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_UndoRedo.cpp', level=3)
doc.add_paragraph('Тесты системы отмены/повтора (17 тестов):')
tests = [
    ('InitiallyEmpty', 'Изначально undo/redo недоступны.'),
    ('CanUndoAfterTwoStates', 'Undo доступен после второго состояния.'),
    ('CanRedoAfterUndo', 'Redo доступен после отмены.'),
    ('UndoRestoresPreviousState', 'Отмена восстанавливает предыдущее состояние.'),
    ('UndoReturnsFalseAtBeginning', 'Undo возвращает false на начальном состоянии.'),
    ('MultipleUndoSteps', 'Последовательные шаги отмены.'),
    ('RedoRestoresNextState', 'Повтор восстанавливает следующее состояние.'),
    ('RedoReturnsFalseAtEnd', 'Redo возвращает false на последнем состоянии.'),
    ('UndoRedoRoundTrip', 'Сложные последовательности undo/redo.'),
    ('DuplicateStateIsNotAdded', 'Дубликаты состояний не добавляются в историю.'),
    ('ClearData', 'Полный сброс истории.'),
    ('NewStateAfterUndoClearsRedo', 'Новое состояние после undo очищает redo-ветку.'),
    ('ActiveIndexPreserved', 'Метаданные (activeIndex) сохраняются при undo/redo.'),
    ('MacroFlagsPreserved', 'Флаги макросов сохраняются при undo/redo.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_page_break()

# Integration
doc.add_heading('7.8. Категория integration', level=2)

doc.add_heading('test_Integration.cpp', level=3)
doc.add_paragraph('Сквозные интеграционные тесты (8 тестов):')
tests = [
    ('ParseTransformVisualize', 'Полный цикл: разбор грамматики → устранение рекурсии → создание диаграммы.'),
    ('LoadFactorizeSave', 'Загрузка → левая факторизация → сохранение.'),
    ('MultipleTransformations', 'Цепочка нескольких преобразований подряд.'),
    ('MultipleNonTerminals', 'Сложная грамматика с множеством нетерминалов.'),
    ('LargeGrammar', 'Стресс-тест с 20 нетерминалами.'),
    ('GrammarCloning', 'Клонирование RE-деревьев через copy().'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

doc.add_heading('test_GrammarPersistence.cpp и test_LoadGrammar.cpp', level=3)
doc.add_paragraph('Тесты сохранения/загрузки:')
tests = [
    ('SaveAndLoad', 'Полный цикл сохранения и загрузки грамматики с правилами.'),
    ('LoadLANGGRM', 'Загрузка реального файла грамматики LANG.GRM с 50+ символами.'),
]
for name, desc in tests:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    p.add_run(f' — {desc}')

# Save
doc.save('/mnt/d/geXrBy/Documents/vkr/Изменения_и_тесты_SynGT.docx')
print('Changes + tests document saved.')
